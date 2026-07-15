#!/usr/bin/env python3
"""按目录.md合并正稿，并生成统一版式的Markdown与DOCX。"""

from __future__ import annotations

import re
import hashlib
import subprocess
from pathlib import Path

from PIL import Image
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TOC = ROOT / "目录.md"
OUT_DIR = ROOT / "成书稿"
OUT_MD = OUT_DIR / "《卡若的IP财富旅程》完整版.md"
OUT_DOCX = OUT_DIR / "《卡若的IP财富旅程》完整版.docx"
FRONT_COVER = ROOT / "封面" / "封面.png"
BACK_COVER = ROOT / "封面" / "尾页.png"
IMAGE_CACHE = Path("/private/tmp/karuo_ip_book_docx_images")

BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 103, 112)
GOLD = RGBColor(166, 113, 38)
INK = RGBColor(30, 35, 40)


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, BLUE),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Part Divider" not in doc.styles:
        part_style = doc.styles.add_style("Part Divider", WD_STYLE_TYPE.PARAGRAPH)
    else:
        part_style = doc.styles["Part Divider"]
    part_style.font.name = "Arial"
    part_style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    part_style.font.size = Pt(24)
    part_style.font.bold = True
    part_style.font.color.rgb = BLUE
    part_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part_style.paragraph_format.space_before = Pt(140)
    part_style.paragraph_format.space_after = Pt(18)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run("《卡若的IP财富旅程》")
    set_run_font(hr, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(footer)
    fr2 = footer.add_run(" —")
    set_run_font(fr2, size=9, color=MUTED)


def parse_toc():
    part = None
    parts: list[tuple[str, list[tuple[str, Path]]]] = []
    for line in TOC.read_text(encoding="utf-8").splitlines():
        if line.startswith("## PART "):
            part = (line[3:].strip(), [])
            parts.append(part)
            continue
        m = re.match(r'- \[([^]]+)\]\(<([^>]+)>\)', line)
        if m and part:
            label, rel = m.groups()
            part[1].append((label, ROOT / rel))
    return parts


def clean_inline(text: str) -> str:
    text = re.sub(r'<video\s+src="([^"]+)"[^>]*></video>', r'视频素材：\1', text)
    text = re.sub(r'\[([^]]+)\]\(([^)]+)\)', r'\1（\2）', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def add_inline_runs(paragraph, text: str):
    chunks = re.split(r'(\*\*.*?\*\*)', text)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(chunk)
            set_run_font(run)


def clean_public_content(content: str) -> str:
    """公开稿删除内部同步链接和原始视频标签，避免出版稿暴露工作流。"""
    content = re.sub(r'<video\s+src="[^"]+"[^>]*></video>\s*', '', content)
    content = re.sub(r'(?m)^\s*https?://[^\s]+\s*$', '', content)
    content = re.sub(r'(?m)^\s*视频素材：[^\n]+\s*$', '', content)
    return content


def image_path_from_line(line: str, source: Path) -> tuple[str, Path | None] | None:
    m = re.fullmatch(r'!\[([^]]*)\]\(([^)]+)\)', line.strip())
    if not m:
        return None
    alt, ref = m.groups()
    ref = ref.replace("%20", " ")
    if ref.startswith("http://") or ref.startswith("https://"):
        return alt, None
    return alt, (source.parent / ref).resolve()


def prepare_docx_image(path: Path) -> tuple[Path, float]:
    """生成不改原图的Word专用压缩副本，并返回合适宽度。"""
    IMAGE_CACHE.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:16]
    cached = IMAGE_CACHE / f"{digest}.jpg"
    if not cached.exists() or cached.stat().st_mtime < path.stat().st_mtime:
        subprocess.run(
            [
                "sips", "-Z", "1600", "-s", "format", "jpeg",
                "-s", "formatOptions", "78", str(path), "--out", str(cached),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
        )
    with Image.open(cached) as img:
        width_px, height_px = img.size
    width_in = min(6.1, 8.0 * width_px / max(height_px, 1))
    return cached, max(width_in, 2.4)


def add_chapter(doc: Document, label: str, source: Path):
    doc.add_page_break()
    text = source.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    p = doc.add_paragraph(style="Heading 1")
    add_inline_runs(p, f"{label.replace(' ', '｜', 1)}")

    title_skipped = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if not title_skipped and line.startswith("# "):
            title_skipped = True
            continue
        if line == "---":
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
            continue
        img = image_path_from_line(line, source)
        if img:
            alt, path = img
            if path and path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    docx_image, width_in = prepare_docx_image(path)
                    shape = run.add_picture(str(docx_image), width=Inches(width_in))
                    shape._inline.docPr.set("descr", alt or path.stem)
                except UnrecognizedImageError:
                    print(f"跳过Word无法识别的图片：{path}")
                    continue
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(10)
                    cr = cap.add_run(alt)
                    set_run_font(cr, size=9, color=MUTED, italic=True)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, clean_inline(line[2:]))
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, clean_inline(line[3:]))
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, clean_inline(line[4:]))
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, clean_inline(line[2:]))
        elif re.match(r"^\d+[）.)]", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, clean_inline(re.sub(r"^\d+[）.)]\s*", "", line)))
        else:
            cleaned = clean_inline(line)
            if cleaned:
                p = doc.add_paragraph()
                add_inline_runs(p, cleaned)


def write_merged_markdown(parts):
    out = [
        "![《卡若的IP财富旅程》封面](../封面/封面.png)",
        "",
        "# 《卡若的IP财富旅程》",
        "",
        "> 一个私域创业者26年的真实经历：从互联网启蒙，到内容、私域与AI的商业进化。",
        "",
        "## 作者序：我为什么要把这些经历写下来",
        "",
        "我不想把这本书写成一套看起来很完整、但离真实生活很远的道理。",
        "",
        "我更愿意把自己走过的路摊开：小时候第一次接触电脑，后来做网站、做游戏、做电商，再到做私域和AI。这里面有赚到钱的时候，也有判断失误、团队离开、公司倒下之后重新开始的时候。",
        "",
        "这些经历没有一条可以直接复制。能留下来的，是我在一次次真实交易、真实合作和真实关系里，慢慢确认的几个问题：客户为什么愿意相信你？流量为什么留不住？一个人怎样借助工具做出过去一个团队的产出？组织如何在分钱、交付和信任之间跑稳？",
        "",
        "所以，这不是成功学，也不是项目说明书。它是一份创业者的长期记录。你可以把它当成故事读，也可以把每一章当成一次复盘：先看现场，再看判断，最后决定哪些动作值得放到自己的生活里试一次。",
        "",
        "## 阅读提示",
        "",
        "书中的时间、项目和数据以现有书稿、项目记录和作者回忆为基础；不同章节的案例证据强度不同，阅读时请把作者经验与普遍规律分开。涉及投资、医疗、法律和平台规则的内容，不构成专业建议。",
        "",
        "## 目录",
        "",
    ]
    chapter_no = 0
    for part_title, chapters in parts:
        out += ["", f"# {part_title}", ""]
        for label, source in chapters:
            chapter_no += 1
            out += [f"## {label.replace(' ', '｜', 1)}", ""]
            content = clean_public_content(source.read_text(encoding="utf-8", errors="ignore"))
            content = re.sub(r"^# .*?\n+", "", content, count=1, flags=re.M)
            for match in re.finditer(r'!\[([^]]*)\]\(([^)]+)\)', content):
                ref = match.group(2)
                if ref.startswith(("http://", "https://")):
                    continue
                absolute = (source.parent / ref.replace("%20", " ")).resolve()
                if absolute.exists():
                    new_ref = absolute.relative_to(ROOT).as_posix()
                    content = content.replace(match.group(0), f"![{match.group(1)}](../{new_ref})")
            out += [content.rstrip(), ""]
    out += ["", "![《卡若的IP财富旅程》尾页](../封面/尾页.png)", ""]
    OUT_MD.write_text("\n".join(out).rstrip() + "\n", encoding="utf-8")


def add_cover_image(doc: Document, image_path: Path, alt: str, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)


def build_docx(parts):
    doc = Document()
    configure_document(doc)

    doc.sections[0].different_first_page_header_footer = True
    add_cover_image(doc, FRONT_COVER, "《卡若的IP财富旅程》封面", width=6.0)

    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    add_inline_runs(p, "作者序：我为什么要把这些经历写下来")
    for paragraph in [
        "我不想把这本书写成一套看起来很完整、但离真实生活很远的道理。",
        "我更愿意把自己走过的路摊开：小时候第一次接触电脑，后来做网站、做游戏、做电商，再到做私域和AI。这里面有赚到钱的时候，也有判断失误、团队离开、公司倒下之后重新开始的时候。",
        "这些经历没有一条可以直接复制。能留下来的，是我在一次次真实交易、真实合作和真实关系里，慢慢确认的几个问题：客户为什么愿意相信你？流量为什么留不住？一个人怎样借助工具做出过去一个团队的产出？组织如何在分钱、交付和信任之间跑稳？",
        "所以，这不是成功学，也不是项目说明书。它是一份创业者的长期记录。你可以把它当成故事读，也可以把每一章当成一次复盘：先看现场，再看判断，最后决定哪些动作值得放到自己的生活里试一次。",
    ]:
        p = doc.add_paragraph()
        add_inline_runs(p, paragraph)
    p = doc.add_paragraph(style="Heading 1")
    add_inline_runs(p, "阅读提示")
    p = doc.add_paragraph()
    add_inline_runs(p, "书中的时间、项目和数据以现有书稿、项目记录和作者回忆为基础；不同章节的案例证据强度不同，阅读时请把作者经验与普遍规律分开。涉及投资、医疗、法律和平台规则的内容，不构成专业建议。")
    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("目录")
    for part_title, chapters in parts:
        p = doc.add_paragraph()
        pr = p.add_run(part_title)
        set_run_font(pr, size=12, color=BLUE, bold=True)
        for label, _ in chapters:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(5)
            add_inline_runs(p, label)

    for part_title, chapters in parts:
        doc.add_page_break()
        p = doc.add_paragraph(style="Part Divider")
        p.add_run(part_title)
        for label, source in chapters:
            add_chapter(doc, label, source)

    back = doc.add_section(WD_SECTION.NEW_PAGE)
    back.top_margin = Inches(0.5)
    back.bottom_margin = Inches(0.5)
    back.left_margin = Inches(0.5)
    back.right_margin = Inches(0.5)
    back.header.is_linked_to_previous = False
    back.footer.is_linked_to_previous = False
    back.header.paragraphs[0].clear()
    back.footer.paragraphs[0].clear()
    add_cover_image(doc, BACK_COVER, "《卡若的IP财富旅程》尾页", width=6.5)

    props = doc.core_properties
    props.title = "卡若的IP财富旅程"
    props.author = "卡若"
    props.subject = "个人IP、私域创业与AI商业"
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main():
    parts = parse_toc()
    missing = [str(path) for _, chapters in parts for _, path in chapters if not path.exists()]
    if missing:
        raise SystemExit("目录存在断链：\n" + "\n".join(missing))
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_merged_markdown(parts)
    build_docx(parts)
    print(f"完成：{sum(len(c) for _, c in parts)}章")
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
